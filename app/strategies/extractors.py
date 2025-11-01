# app/strategies/extractors.py
from docling.document_converter import DocumentConverter
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document
from typing import List
from .base import ExtractorStrategy
from pypdf import PdfReader
import os
import re
from docx import Document as DocxDocument
# docling
class DoclingExtractor(ExtractorStrategy):
    def extract(self, file_path: str) -> List[Document]:
        converter = DocumentConverter()
        result = converter.convert(file_path)
        return [Document(page_content=result.document.export_to_markdown(), metadata={"page": 0})]
# pypdf
class LangChainPyMuPDFExtractor(ExtractorStrategy):
    def extract(self, file_path: str) -> List[Document]:
        loader = PyMuPDFLoader(file_path)
        return loader.load()

#pdfFA
class PyPDFExtractor(ExtractorStrategy):
    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,!?؟٪%\u0600-\u06FF]', '', text)
        return text.strip()

    def extract(self, file_path: str) -> List[Document]:
        docs = []
        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, start=1):
                raw_text = page.extract_text()
                if raw_text and raw_text.strip():
                    cleaned_text = self._clean_text(raw_text)
                    docs.append(Document(
                        page_content=cleaned_text,
                        metadata={"source": os.path.basename(file_path), "page": page_num}
                    ))
        except Exception as e:
            print(f"Error processing PDF with pypdf: {e}")
        return docs
# docx


class DocxExtractor(ExtractorStrategy):
    def extract(self, file_path: str) -> List[Document]:
        """
        Extracts the entire DOCX file as a single 'page' document.
        The page number is always set to 1 to avoid issues with chunking.
        """
        docs = []
        try:
            docx = DocxDocument(file_path)
            all_paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
            if not all_paragraphs:
                return []

            page_content = "\n\n".join(all_paragraphs)
            docs.append(Document(
                page_content=page_content,
                metadata={
                    "source": os.path.basename(file_path),
                    "page": 1,       
                    "type": "docx"
                }
            ))

        except Exception as e:
            print(f"Error processing DOCX: {e}") 
        
        return docs
